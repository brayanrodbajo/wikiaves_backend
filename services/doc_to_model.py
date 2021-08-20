from django.core.exceptions import ObjectDoesNotExist
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings

from birds.models import Bird, Family, Order
from birds.serializers import BirdSerializer, FamilySerializer, OrderSerializer

from docx import Document
from os.path import isfile
import unidecode
import requests
import json
import re

headers = {
    'Authorization': 'Token 75c0fa6dd6987219a1ec4046ed8ba710bd81069f',
    'Content-Type': 'application/json',
}
url = 'http://i2thub.icesi.edu.co/wikiaves_dev/api/birds/'
XCurl = 'https://www.xeno-canto.org/api/2/recordings?query=nr:'


def handle_uploaded_file(f):
    path = 'tmp/ficha.'+f.name.split('.', 1)[1]
    return default_storage.save(path, ContentFile(f.read()))


def delete_uploaded_file(path):
    default_storage.delete(path)


def tag_html_para(para):
    result = ""
    for run in para.runs:
        if run.bold and run.italic:
            result += "<b><i>" + run.text + "</b></i>"
        elif run.bold:
            result += "<b>" + run.text + "</b>"
        elif run.italic:
            result += "<i>" + run.text + "</i>"
        else:
            result += run.text
    else:
        return result


def list_text_paras(index, all_paras):
    result = [tag_html_para(all_paras[index])]
    while len(all_paras[index + 1].text.strip()) > 0:
        result.append(tag_html_para(all_paras[index + 1]))
        index += 1
    return result


def list_paras(index, all_paras):
    result = [all_paras[index]]
    while len(all_paras[index + 1].text.strip()) > 0:
        result.append(all_paras[index + 1])
        index += 1
    return result


def join_text_paras(index, all_paras):
    result = tag_html_para(all_paras[index])
    while len(all_paras[index + 1].text.strip()) > 0:
        result += "<br>" + tag_html_para(all_paras[index + 1])
        index += 1
    return result


def list_not_num_paras(index, all_paras):
    result = [all_paras[index]]
    while (index < len(all_paras) - 1 and all_paras[index + 1]._p.pPr.numPr is None):
        result.append(all_paras[index + 1])
        index += 1
    return result


def bird_exists(file_path):
    try:
        if isfile(file_path):
            document = Document(file_path)
            all_paras = document.paragraphs
            all_paras_txt = [para.text for para in all_paras]
            first_line = all_paras_txt[0]
            sci_name = first_line.split('(')[0].strip()
            # If bird exists
            try:
                bird = Bird.objects.get(scientific_names__name=sci_name)
                return {
                    "success": True,
                    "exists": True,
                    "bird_id": bird.id,
                    "scientific_name": sci_name
                }
            except ObjectDoesNotExist:
                return {
                    "success": True,
                    "exists": False
                }
    except Exception as e:
        return {
            "success": False,
            "msg": str(e),
            "error_code": "Exception",
        }


def doc_to_model(file_path, author_id=None):
    log = "Comienzo"
    try:
        if isfile(file_path):
            payload = {}
            document = Document(file_path)
            all_paras = document.paragraphs
            log = "Párrafos"
            all_paras_txt = [para.text for para in all_paras]
            first_line = all_paras_txt[0]
            log = "Nombre científico"
            sci_name = first_line.split('(')[0].strip()
            payload['scientific_names'] = [{
                "name": sci_name,
                "main": True
            }]
            log = "Nombres comunes"
            if '(' in first_line:
                common_names = first_line[first_line.find("(") + 1:first_line.find(")")]
                lan = 'en'
                payload['common_names'] = []
                for common_names_lan in common_names.split('|'):
                    main = True
                    for common_name in common_names_lan.split('/'):
                        payload['common_names'].append({
                            "name": {
                                "language": lan,
                                "text": common_name.strip()
                            },
                            "main": main
                        })
                        main = False
                    lan = 'es'
            order_id = None
            payload_family = None
            log = "Orden y familia"
            for para in all_paras_txt:
                if para.startswith('Orden'):
                    order = para[6:].strip()
                    try:
                        order_obj = Order.objects.get(scientific_names__name=order)
                        order_id = order_obj.id
                    except ObjectDoesNotExist:
                        payload_order = {
                            "scientific_names": [
                                {
                                    "name": order,
                                    "main": True
                                }
                            ]
                        }
                        serializer = OrderSerializer(data=payload_order)
                        if serializer.is_valid():
                            obj = serializer.save()
                            order_id = obj.id
                        else:
                            return {
                                "success": False,
                                "msg": serializer.errors,
                                "section": "Order creation"
                            }
                        if payload_family:
                            response = requests.request("POST", url + 'families', headers=headers,
                                                        data=json.dumps(payload_family))
                            family_id = json.loads(response.content)['id']
                            payload['family'] = family_id
                elif para.startswith('Familia'):
                    family = para[8:].strip()
                    try:
                        family_obj = Family.objects.get(scientific_names__name=family)
                        payload['family'] = family_obj.id
                    except ObjectDoesNotExist:
                        payload_family = {
                            "scientific_names": [
                                {
                                    "name": family,
                                    "main": True
                                }
                            ],
                            "order": order_id
                        }
                        if order_id:
                            serializer = FamilySerializer(data=payload_family)
                            if serializer.is_valid():
                                obj = serializer.save()
                                family_id = obj.id
                                payload['family'] = family_id
                            else:
                                return {
                                    "success": False,
                                    "msg": serializer.errors,
                                    "section": "Family creation"
                                }
            all_paras_no_acc = [unidecode.unidecode(para.strip()) for para in all_paras_txt]  # without accent
            if 'Descripcion' in all_paras_no_acc or 'Descripcion:' in all_paras_no_acc:
                log = "Descripción"
                if 'Descripcion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Descripcion')
                else:
                    index = all_paras_no_acc.index('Descripcion:')
                desc = join_text_paras(index + 1, all_paras)
                payload['description'] = {
                    'language': 'es',
                    'text': '<p>' + desc + '</p>'
                }
            if 'Distribucion' in all_paras_no_acc or 'Distribucion:' in all_paras_no_acc:
                log = "Distribución"
                if 'Distribucion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Distribucion')
                else:
                    index = all_paras_no_acc.index('Distribucion:')
                dist = join_text_paras(index + 1, all_paras)
                payload['distribution'] = {
                    'text': {
                        'language': 'es',
                        'text': '<p>' + dist + '</p>'
                    }
                }
            if 'Subespecies' in all_paras_no_acc or 'Subespecies:' in all_paras_no_acc:
                log = "Subespecies"
                if 'Subespecies' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Subespecies')
                else:
                    index = all_paras_no_acc.index('Subespecies:')
                subs = list_paras(index + 1, all_paras)
                payload['subspecies'] = []
                for para in subs:
                    name = para.text.split(':')[0].strip()
                    dist = tag_html_para(para).split(':', 1)[1].strip()[8:]
                    payload['subspecies'].append({
                        'names': [{
                            'name': name,
                            'main': True
                        }],
                        'distribution': {
                            'text': {
                                'language': 'es',
                                'text': '<p>' + dist + '</p>'
                            }
                        },
                        'lengths': [],
                        'weights': []
                    })
            if 'Habitat' in all_paras_no_acc or 'Habitat:' in all_paras_no_acc:
                log = "Hábitat"
                if 'Habitat' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Habitat')
                else:
                    index = all_paras_no_acc.index('Habitat:')
                hab = join_text_paras(index + 1, all_paras)
                payload['habitat'] = {
                    'language': 'es',
                    'text': '<p>' + hab + '</p>'
                }
            if 'Identificacion' in all_paras_no_acc or 'Identificacion:' in all_paras_no_acc:
                log = "Identificacion"
                if 'Identificacion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Identificacion')
                else:
                    index = all_paras_no_acc.index('Identificacion:')
                list_dist = list_paras(index + 1, all_paras)
                payload['identification'] = {}
                for ind, para in enumerate(list_dist):
                    if para.text.startswith('Descripcion') or para.text.startswith('Descripción'):
                        log = "Identificacion - Descripción"
                        desc = tag_html_para(para)[13:]
                        payload['identification']['description'] = {
                            'language': 'es',
                            'text': '<p>' + desc + '</p>'
                        }
                    if para.text.startswith('Medidas'):
                        log = "Identificacion - Medidas"
                        list_med = list_not_num_paras(ind + 1, list_dist)
                        for med in list_med:
                            subspecies = None
                            name = med.text.split(':')[0]
                            value = med.text.split(':', 1)[1].strip()
                            if name.startswith('Longitud'):
                                measure = 'lengths'
                            elif name.startswith('Peso'):
                                measure = 'weights'
                            if '(' in name:  # It has a subspecies
                                subspecies = name[name.find("(") + 1:name.find(")")].strip()
                                name = name.split('(')[0].strip()
                                subspecies_obj = next(item for item in payload['subspecies'] if
                                                      subspecies in item['names'][0]['name'])
                            if not measure in payload['identification']:
                                payload['identification'][measure] = []
                            ref = None
                            if not ',' in value:
                                if "(" in value:
                                    ref_idx_start = value.index('(')
                                    ref = value[ref_idx_start:].replace('.', '').strip()
                                    cond_1 = '-' in value[:ref_idx_start]
                                    cond_2 = '–' in value[:ref_idx_start]
                                else:
                                    cond_1 = '-' in value
                                    cond_2 = '–' in value
                                if cond_1:
                                    inferior = float(value.split('-')[0].replace(' ', ''))
                                    superior_text = re.findall("\d+[.]?\d*", value.split('-')[1])[0]
                                    superior = float(superior_text.replace(' ', ''))
                                    unit_idx = value.index(superior_text) + len(superior_text)
                                    unit = re.split('\s|\.|\(', value[unit_idx:].strip())[0]
                                    if subspecies:
                                        subspecies_obj[measure] += [{
                                            'name': name,
                                            'value': {
                                                'inferior': inferior,
                                                'superior': superior
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                                    else:
                                        payload['identification'][measure] += [{
                                            'name': name,
                                            'value': {
                                                'inferior': inferior,
                                                'superior': superior
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                                elif cond_2:
                                    inferior = float(value.split('–')[0].replace(' ', ''))
                                    superior_text = re.findall("\d+[.]?\d*", value.split('–')[1])[0]
                                    superior = float(superior_text.replace(' ', ''))
                                    unit_idx = value.index(superior_text) + len(superior_text)
                                    unit = re.split('\s|\.', value[unit_idx:].strip())[0]
                                    if subspecies:
                                        subspecies_obj[measure] += [{
                                            'name': name,
                                            'value': {
                                                'inferior': inferior,
                                                'superior': superior
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                                    else:
                                        payload['identification'][measure] += [{
                                            'name': name,
                                            'value': {
                                                'inferior': inferior,
                                                'superior': superior
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                                else:
                                    average_txt = re.findall("\d+[.]?\d*", value)[0]
                                    average = float(average_txt.replace(' ', ''))
                                    unit_idx = value.index(average_txt) + len(average_txt)
                                    unit = re.split('\s|\.|\(', value[unit_idx:].strip())[0]
                                    if subspecies:
                                        subspecies_obj[measure] += [{
                                            'name': name,
                                            'value': {
                                                'average': average
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                                    else:
                                        payload['identification'][measure] += [{
                                            'name': name,
                                            'value': {
                                                'average': average
                                            },
                                            'unit': unit,
                                            'reference': ref
                                        }]
                if payload['identification'] == {}:
                    del payload['identification']
            if 'Migracion' in all_paras_no_acc or 'Migracion:' in all_paras_no_acc:
                log = "Migración"
                if 'Migracion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Migracion')
                else:
                    index = all_paras_no_acc.index('Migracion:')
                mig = join_text_paras(index + 1, all_paras)
                name = None
                if '.' in mig.strip()[:-1]:
                    mig_type = mig.split('.')[0].strip()
                    mig = mig.split('.', 1)[1].strip()
                    name = {
                        "language": "es",
                        "text": mig_type
                    }
                payload['migration'] = {
                    "name": name,
                    "text": {
                        "language": "es",
                        "text": '<p>' + mig + '</p>'
                    }
                }
            if 'Alimentacion' in all_paras_no_acc or 'Alimentacion:' in all_paras_no_acc:
                log = "Alimentacion"
                if 'Alimentacion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Alimentacion')
                else:
                    index = all_paras_no_acc.index('Alimentacion:')
                alim = join_text_paras(index + 1, all_paras)
                payload['feeding'] = {
                    "text": {
                        "language": "es",
                        "text": '<p>' + alim + '</p>'
                    }
                }
            payload['vocalizations'] = []
            audio_cantos = []
            audio_llamados = []
            for para in all_paras_txt[:6]:
                if para.startswith('Canto'):
                    log = "Canto"
                    canto_ids = para[5:].replace(':', '').strip()
                    main = True
                    for canto_id in canto_ids.split('/'):
                        canto_id = canto_id.strip()
                        if canto_id.startswith('XC'):
                            canto_id_n = canto_id[2:]
                            response = requests.request("GET", XCurl + canto_id_n)
                            response_dict = json.loads(response.content)['recordings']
                            if len(response_dict) > 0:
                                audio_cantos.append({
                                    'url': 'https:' + response_dict[0]['url'],
                                    'XC_id': canto_id,
                                    'main': main
                                })
                            main = False
                if para.startswith('Llamado'):
                    log = "Llamado"
                    llamado_ids = para[7:].replace(':', '').strip()
                    main = True
                    for llamado_id in llamado_ids.split('/'):
                        llamado_id = llamado_id.strip()
                        if llamado_id.startswith('XC'):
                            llamado_id_n = llamado_id[2:]
                            response = requests.request("GET", XCurl + llamado_id_n)
                            response_dict = json.loads(response.content)['recordings']
                            if len(response_dict) > 0:
                                audio_llamados.append({
                                    'url': 'https:' + response_dict[0]['url'],
                                    'XC_id': llamado_id,
                                    'main': main
                                })
                            main = False
            if 'Vocalizaciones' in all_paras_no_acc or 'Vocalizaciones:' in all_paras_no_acc:
                log = "Vocalizaciones"
                if 'Vocalizaciones' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Vocalizaciones')
                else:
                    index = all_paras_no_acc.index('Vocalizaciones:')
                list_voc = list_text_paras(index + 1, all_paras)
                for para in list_voc:
                    if para.startswith('Canto'):
                        log = "Vocalizaciones - Canto"
                        if len(audio_cantos) > 0:
                            payload['vocalizations'].append({
                                'category': 'SONG',
                                'long_description': {
                                    "language": "es",
                                    "text": '<p>' + para[7:] + '</p>'
                                },
                                'xenocantos': audio_cantos
                            })
                        else:
                            payload['vocalizations'].append({
                                'category': 'SONG',
                                'long_description': {
                                    "language": "es",
                                    "text": '<p>' + para[7:] + '</p>'
                                }
                            })
                    elif para.startswith('Llamado'):
                        log = "Vocalizaciones - Llamado"
                        if len(audio_llamados) > 0:
                            payload['vocalizations'].append({
                                'category': 'CALL',
                                'long_description': {
                                    "language": "es",
                                    "text": '<p>' + para[9:] + '</p>'
                                },
                                'xenocantos': audio_llamados
                            })
                        else:
                            payload['vocalizations'].append({
                                'category': 'CALL',
                                'long_description': {
                                    "language": "es",
                                    "text": '<p>' + para[9:] + '</p>'
                                }
                            })
                if len(payload['vocalizations']) == 1:
                    if payload['vocalizations'][0]['category'] == 'SONG':
                        if len(audio_llamados) > 0:
                            payload['vocalizations'].append({
                                'category': 'CALL',
                                'xenocantos': audio_llamados
                            })
                    else:
                        if len(audio_cantos) > 0:
                            payload['vocalizations'].append({
                                'category': 'SONG',
                                'xenocantos': audio_cantos
                            })
                elif len(payload['vocalizations']) == 0:
                    if len(audio_cantos) > 0:
                        payload['vocalizations'].append({
                            'category': 'SONG',
                            'xenocantos': audio_cantos
                        })
                    if len(audio_llamados) > 0:
                        payload['vocalizations'].append({
                            'category': 'CALL',
                            'xenocantos': audio_llamados
                        })
            else:
                if len(audio_cantos) > 0:
                    payload['vocalizations'].append({
                        'category': 'SONG',
                        'xenocantos': audio_cantos
                    })
                if len(audio_llamados) > 0:
                    payload['vocalizations'].append({
                        'category': 'CALL',
                        'xenocantos': audio_llamados
                    })
            if 'Reproduccion' in all_paras_no_acc or 'Reproduccion:' in all_paras_no_acc:
                log = "Reproducción"
                if 'Reproduccion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Reproduccion')
                else:
                    index = all_paras_no_acc.index('Reproduccion:')
                if all_paras[index + 1]._p.pPr.numPr is None:  # No categories
                    repr = join_text_paras(index + 1, all_paras)
                    payload['reproduction'] = [{
                        "name": {
                            "language": "es",
                            "text": "Información General"
                        },
                        "text": {
                            "language": "es",
                            "text": '<p>' + repr + '</p>'
                        }
                    }]
                else:
                    list_repr = list_text_paras(index + 1, all_paras)
                    if not ':' in list_repr[0]:  # (!!!!) MOMENTÁNEAMENTE
                        repr = join_text_paras(index + 1, all_paras)
                        payload['reproduction'] = [{
                            "name": {
                                "language": "es",
                                "text": "Información General"
                            },
                            "text": {
                                "language": "es",
                                "text": '<p>' + repr + '</p>'
                            }
                        }]
                    else:
                        payload['reproduction'] = []
                        for para in list_repr:
                            r_type = para.split(':')[0].strip()
                            r_txt = para.split(':', 1)[1].strip()
                            payload['reproduction'].append({
                                "name": {
                                    "language": "es",
                                    "text": r_type
                                },
                                "text": {
                                    "language": "es",
                                    "text": '<p>' + r_txt + '</p>'
                                }
                            })
                        if len(payload['reproduction']) == 0:
                            del payload['reproduction']
            if 'Comportamiento' in all_paras_no_acc or 'Comportamiento:' in all_paras_no_acc:
                log = "Comportamiento"
                if 'Comportamiento' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Comportamiento')
                else:
                    index = all_paras_no_acc.index('Comportamiento:')
                list_comp = list_text_paras(index + 1, all_paras)
                payload['behavior'] = []
                for para in list_comp:
                    if ':' in para:
                        c_type = para.split(':')[0].strip()
                        c_txt = para.split(':', 1)[1].strip()
                    else:
                        c_type = para.split('.')[0].strip()
                        c_txt = para.split('.', 1)[1].strip()
                    payload['behavior'].append({
                        "name": {
                            "language": "es",
                            "text": c_type
                        },
                        "text": {
                            "language": "es",
                            "text": '<p>' + c_txt + '</p>'
                        }
                    })
            if 'Conservacion' in all_paras_no_acc or 'Conservacion:' in all_paras_no_acc:
                log = "Conservacion"
                if 'Conservacion' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Conservacion')
                else:
                    index = all_paras_no_acc.index('Conservacion:')
                conserva = join_text_paras(index + 1, all_paras)
                conserva_type = conserva.split('.')[0].strip()
                conserva = conserva.split('.', 1)[1].strip()
                payload['conservation'] = {
                    "name": {
                        "language": "es",
                        "text": conserva_type
                    },
                    "text": {
                        "language": "es",
                        "text": '<p>' + conserva + '</p>'
                    }
                }
            if 'Taxonomia' in all_paras_no_acc or 'Taxonomia:' in all_paras_no_acc:
                log = "Taxonomia"
                if 'Taxonomia' in all_paras_no_acc:
                    index = all_paras_no_acc.index('Taxonomia')
                else:
                    index = all_paras_no_acc.index('Taxonomia:')
                taxono = join_text_paras(index + 1, all_paras)
                payload['taxonomy'] = {
                    'language': 'es',
                    'text': '<p>' + taxono + '</p>'
                }
            # If bird exists
            log = "Bird Complete"
            try:
                bird_obj = Bird.objects.get(scientific_names__name=sci_name)
                payload['authors'] = [author_id]
                serializer = BirdSerializer(bird_obj, data=payload)
                if serializer.is_valid():
                    serializer.save()
                    bird_dict = serializer.data
                else:
                    return {
                        "success": False,
                        "msg": serializer.errors
                    }
            except ObjectDoesNotExist:
                serializer = BirdSerializer(data=payload)
                if serializer.is_valid():
                    serializer.save()
                    bird_dict = serializer.data
                else:
                    return {
                        "success": False,
                        "msg": serializer.errors
                    }
            return {
                "success": True,
                "bird": bird_dict
            }
    except Exception as e:
        return {
            "success": False,
            "msg": str(e),
            "exception": True,
            "section": log
        }
